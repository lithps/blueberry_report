import pandas as pd
import yfinance as yf
import matplotlib.pyplot as plt
from pytz import timezone
import datetime
import matplotlib.pyplot as plt
from datetime import date
import os
from docx import Document
from docx.shared import Inches
#
# Function to fetch stock data
def fetch_stock_data(ticker, start_date=None, end_date=None, period=None):
    if period:
        return yf.download(ticker, period=period)[['Open', 'Adj Close', 'Volume']]
    else:
        return yf.download(ticker, start=start_date, end=None)[['Open', 'Adj Close', 'Volume']]

# Fetch LINE stock info
ticker = 'LINE'
LINE_stock_data = fetch_stock_data(ticker, period='max')
LINE_stock_data.index = LINE_stock_data.index.tz_localize(None)
start_date = LINE_stock_data.index.min()

# Ensure the datetime index is timezone-naive
LINE_stock_data.columns = ['Open', 'Adj Close', 'Volume']

# Fetch peer stock info
peer_data = pd.read_csv('/home/josh/blueberry_report/peers.csv')
peer_tickers = peer_data['Ticker'].tolist()

# Initialize an empty list to store dataframes
dataframes = []

# Loop through each ticker to download and rename the data
for ticker in peer_tickers:
    # Download the stock data
    ticker_data = fetch_stock_data(ticker, start_date=start_date, end_date=None)
    
    # Rename the columns
    ticker_data.columns = [f'{ticker}_{col}' for col in ['Open', 'Adj Close', 'Volume']]
    
    # Append the dataframe to the list
    dataframes.append(ticker_data)

# Concatenate peer dataframes
combined_peer_data = pd.concat(dataframes, axis=1)
combined_peer_data.index = combined_peer_data.index.tz_localize(None)

# Combine LINE ticker with peer tickers
all_tickers = peer_tickers + ['LINE']
unique_tickers = set(all_tickers)

# Initialize lists to store shares outstanding and earnings dates
shares_outstanding_list = []
earnings_dates_list = []
net_income_list = []

# Fetch shares outstanding and earnings dates for each ticker
for t in unique_tickers:
    stock = yf.Ticker(t)
    shares_outstanding = stock.info.get('sharesOutstanding', None)
    net_income = stock.financials.loc['Net Income'].iloc[0] if not stock.financials.empty else None
    
    # Fetch historical earnings dates
    earnings_dates = stock.earnings_dates
    if not earnings_dates.empty:
        # Convert pd.Timestamp.today() to a timezone-aware datetime
        tz = timezone('America/New_York')
        today_tz_aware = pd.Timestamp.now(tz)

        # Filter out future dates
        past_earnings_dates = earnings_dates[earnings_dates.index <= today_tz_aware]
        latest_earnings_date = past_earnings_dates.index[0].date() if not past_earnings_dates.empty else None
    else:
        latest_earnings_date = None
    
    shares_outstanding_list.append({
        'Ticker': t,
        'Latest Earnings Date': latest_earnings_date,
        'Shares Outstanding': shares_outstanding,
        'Net Income': net_income
    })

# Create DataFrame for shares outstanding
financials_df = pd.DataFrame(shares_outstanding_list)
financials_df.set_index('Ticker', inplace=True)
financials_df.index = financials_df.index.astype(str)

# Rename columns
LINE_stock_data.columns = [f'LINE_{col}' for col in LINE_stock_data.columns]


# Concatenate the dataframes
combined_data = pd.concat([LINE_stock_data, combined_peer_data], axis=1)

for ticker in unique_tickers:
    # Calculate vol %
    vol_col = f'{ticker}_Volume'
    vol_per_col = f'{ticker}_Vol_%'
    if vol_col in combined_data:
        combined_data[vol_per_col] = (combined_data[vol_col]/financials_df.loc[ticker, 'Shares Outstanding']) * 100

# Loop through each ticker to calculate EPS
for ticker in unique_tickers:
    eps_col = f'{ticker}_EPS'
    if ticker in financials_df.index:
        combined_data[eps_col] = financials_df.loc[ticker, 'Net Income'] / financials_df.loc[ticker, 'Shares Outstanding']

# Loop through each ticker to calculate P/E ratio
for ticker in unique_tickers:
    close_col = f'{ticker}_Adj Close'
    eps_col = f'{ticker}_EPS'
    PE_col = f'{ticker}_P/E_Ratio'
    if eps_col in combined_data.columns:
        combined_data[PE_col] = (combined_data[close_col] / combined_data[eps_col])

# Filter out NaN and print the updated combined_data DataFrame
filtered_combined_data = combined_data.dropna()

# Grab today's date and create folder for report
today = date.today()
folder_name = today.strftime("%Y-%m-%d")
report_path = os.path.join('/home/josh/blueberry_report/finished_reports', folder_name)
os.makedirs(report_path, exist_ok=True)

# Load the CSV file into a DataFrame
df = filtered_combined_data

# Create volume % df
vol_percent = df.filter(like='Vol_%')
# EPS ratio df
eps = df.filter(like='EPS')
eps = eps.loc[eps.index.max()]
# P/E ratio df
pe_ratio = df.filter(like='P/E')
# Stock price df
stock_price = df.filter(like='Close')
# Daily return
daily_return = stock_price.pct_change().dropna()

# Plot daily returns
daily_return_image = f'{report_path}/daily_returns.png'
plt.figure(figsize=(12, 8))
for column in daily_return.columns:
    if 'LINE_Adj Close' in column:
        plt.plot(daily_return.index, daily_return[column], label=column, color='green', linewidth=2, alpha=0.9)
    else:
        plt.plot(daily_return.index, daily_return[column], label=column, color='black', linewidth=1, alpha=0.5)
plt.title('Daily Returns Post-IPO', fontsize=16)
plt.xlabel('Date', fontsize=14)
plt.ylabel('Daily Return', fontsize=14)
plt.grid(True, alpha=0.3)
if plt.gca().get_legend_handles_labels()[1]:  # Check if there are labels for the legend
    plt.legend()
plt.savefig(daily_return_image)
plt.show()

# Plot volume %
vol_image = f'{report_path}/vol_per.png'
plt.figure(figsize=(12, 8))
for column in vol_percent.columns:
    if 'LINE_Vol_%' in column:
        plt.plot(vol_percent.index, vol_percent[column], label=column, color='green', linewidth=2, alpha=0.9)
    else:
        plt.plot(vol_percent.index, vol_percent[column], label=column, color='black', linewidth=1, alpha=0.5)
plt.title('Volume as Percentage of Shares Outstanding', fontsize=16)
plt.xlabel('Date', fontsize=14)
plt.ylabel('Volume %', fontsize=14)
plt.grid(True, alpha=0.3)
if plt.gca().get_legend_handles_labels()[1]:  # Check if there are labels for the legend
    plt.legend()
plt.savefig(vol_image)
plt.show()

# Bar EPS comparison
eps_image = f'{report_path}/eps.png'
plt.figure(figsize=(12, 8))
eps_tickers = eps.index
eps_values = eps.values
eps_colors = plt.cm.tab20.colors
bars = plt.bar(eps_tickers, eps_values, color=eps_colors[:len(eps_tickers)], label=eps_tickers)
plt.title('EPS Comparison', fontsize=16)
plt.legend(title='Stock Ticker', loc='upper left')
plt.xticks([])
plt.ylabel('EPS Values', fontsize=14)
plt.savefig(eps_image)
plt.show()

# Plot P/E ratio
pe_image = None
if eps['LINE_EPS'] > 0:
    pe_image = f'{report_path}/pe_ratio.png'
    plt.figure(figsize=(12, 8))
    for column in pe_ratio.columns:
        if 'LINE_P/E_Ratio' in column:
            plt.plot(pe_ratio.index, pe_ratio[column], label=column, color='green', linewidth=2, alpha=0.9)
        else:
            plt.plot(pe_ratio.index, pe_ratio[column], label=column, color='black', linewidth=1, alpha=0.5)
    plt.title('P/E Ratio Compared to Peers', fontsize=16)
    plt.xlabel('Date', fontsize=14)
    plt.ylabel('P/E Ratio', fontsize=14)
    plt.grid(True, alpha=0.3)
    if plt.gca().get_legend_handles_labels()[1]:  # Check if there are labels for the legend
        plt.legend()
    plt.savefig(pe_image)
    plt.show()
else:
    pe_image = None

print("Plots generated and saved successfully.")

# Define the template tags with correct image paths
template_tags = {
    '<<Today>>': f'{folder_name}',
    '<<Fig 1 Name>>': 'Daily Returns',
    '<<Fig 1 Description>>': 'Daily Returns are used to measure the volatility and trajectory of a stock. When compared to peers it can help evaluate the performance of one stock versus others. Daily Returns are calculated by taking the difference between the current price and the previous day\'s price and dividing it by the previous day\'s price. The green line below shows LINE\'s daily returns compared to its peer group.',
    '<<Fig 1 Image>>': daily_return_image,
    '<<Fig 2 Name>>': 'Volume as a Percentage of Shares Outstanding',
    '<<Fig 2 Description>>': 'Trading volume can be used to indicate how liquid a company\'s stock is. Excessive trading volumes can show either positive or negative market sentiment depending on the stock trajectory. Trading volume as a percentage of shares outstanding can help a company determine if there is more or less interest in their stock (for good or bad) on a given day.',
    '<<Fig 2 Image>>': vol_image,
    '<<Fig 3 Name>>': 'Earnings per Share Comparison',
    '<<Fig 3 Description>>': 'Earnings per Share (EPS) measures a company\'s profitability. It\'s calculated by subtracting preferred dividends from Net Income and dividing that value by the number of outstanding shares. The higher EPS a company has, the more profitable they are. This calculation is a simplified calculation that takes the Net Income and divides it by the number of outstanding shares.',
    '<<Fig 3 Image>>': eps_image,
}

# Add P/E ratio tag only if pe_image is not None
if pe_image:
    template_tags.update({
        '<<Fig 4 Name>>': 'Price-to-Earnings Ratio Comparison',
        '<<Fig 4 Description>>': 'Price-to-Earnings Ratios measures a company\'s share price relative to their EPS. Lower P/E ratios tend to indicate undervalued companies, whereas a higher P/E ratio tends to indicate a company is overvalued or investors anticipate high growth rates.',
        '<<Fig 4 Image>>': pe_image
    })

# Load the Word document
template = Document('/home/josh/blueberry_report/template_report.docx')

# Replace text tags in the document
for tag, value in template_tags.items():
    if 'Image' not in tag:
        for paragraph in template.paragraphs:
            if tag in paragraph.text:
                paragraph.text = paragraph.text.replace(tag, value)

# Insert images
for tag, image_path in template_tags.items():
    if 'Image' in tag and image_path is not None and isinstance(image_path, str) and os.path.exists(image_path):
        for paragraph in template.paragraphs:
            if tag in paragraph.text:
                paragraph.text = paragraph.text.replace(tag, '')
                run = paragraph.add_run()
                try:
                    run.add_picture(image_path, width=Inches(6))
                except Exception as e:
                    print(f"Error inserting image {image_path}: {e}")

# Remove Figure 4 tags if no image was created
if not pe_image:
    for paragraph in template.paragraphs:
        if '<<Fig 4 Name>>' in paragraph.text or '<<Fig 4 Description>>' in paragraph.text or '<<Fig 4 Image>>' in paragraph.text:
            paragraph.clear()

# Save the updated document
template.save(f'{report_path}/LINE_Stock_Report_{folder_name}.docx')

print("Document updated successfully.")
